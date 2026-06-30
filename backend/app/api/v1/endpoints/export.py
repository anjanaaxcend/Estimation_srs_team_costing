import io
import zipfile
from datetime import datetime
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from fastapi import APIRouter, Response, Body, Depends
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from app.schemas.srs import SRSGenerationResult
from app.core.database import get_db
from app.api.deps import get_optional_user
from app.models.user import User, UserHistory
from app.schemas.cost import CostEstimationExportRequest
from app.services.planning_sync import (
    build_module_timeline,
    build_synced_team_section,
    calculate_cost_totals,
    determine_project_weeks,
    safe_float,
    is_testing_role,
    is_deployment_role,
    is_management_role,
)

router = APIRouter()

def _make_styles():
    """Unified professional styling for Estimator AI Excel exports."""
    # Palette format strictly flat B&W
    BLACK = "000000"
    WHITE = "FFFFFF"

    title_font = Font(name="Calibri", size=18, bold=True, color=BLACK)
    header_font = Font(name="Calibri", size=11, bold=True, color=BLACK)
    bold_font = Font(name="Calibri", size=11, bold=True, color=BLACK)
    normal_font = Font(name="Calibri", size=10, color=BLACK)
    
    white_fill = PatternFill(start_color=WHITE, end_color=WHITE, fill_type="solid")
    
    thin_side = Side(border_style="thin", color=BLACK)
    border = Border(top=thin_side, left=thin_side, right=thin_side, bottom=thin_side)
    
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    return {
        "title_font": title_font,
        "header_font": header_font,
        "bold_font": bold_font,
        "normal_font": normal_font,
        "navy_fill": white_fill,
        "slate_fill": white_fill,
        "alt_fill": white_fill,
        "border": border,
        "center": center_align,
        "left": left_align,
        "accent_color": BLACK
    }

def _write_header_row(ws, row, columns, s, fill=None):
    if fill is None:
        fill = s["slate_fill"]
    for col, text in enumerate(columns, 1):
        c = ws.cell(row=row, column=col, value=text)
        c.font = s["header_font"]
        c.fill = fill
        c.border = s["border"]
        c.alignment = s["center"]

def _write_data_row(ws, row, values, s, alt=False, bold_first=False):
    for col, val in enumerate(values, 1):
        c = ws.cell(row=row, column=col, value=val)
        c.font = s["bold_font"] if (bold_first and col == 1) else s["normal_font"]
        c.border = s["border"]
        c.alignment = s["left"]
        if alt:
            c.fill = s["alt_fill"]

@router.post("/excel")
async def export_excel(
    data: SRSGenerationResult,
    db: Session = Depends(get_db),
    current_user: User | None = Depends(get_optional_user)
):
    if current_user:
        history = UserHistory(
            user_id=current_user.id,
            action="Downloaded SRS Excel",
            details=f"Exported SRS Excel: {data.title}"
        )
        db.add(history)
        db.commit()

    output = io.BytesIO()
    wb = Workbook()
    s = _make_styles()

    # ── Sheet 1: Project Overview ─────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Project Overview"
    ws1.merge_cells("A1:D1")
    ws1["A1"] = f"PROJECT SCOPE: {data.title}"
    ws1["A1"].font = s["title_font"]
    ws1["A1"].fill = s["navy_fill"]
    ws1["A1"].alignment = s["center"]
    ws1.row_dimensions[1].height = 40

    req = data.structured_requirements
    if req:
        rows = [
            ("Project Name", req.project_name or "Unnamed Project"),
            ("Executive Summary", req.executive_summary or ""),
            ("Problem Statement", req.problem_statement or ""),
            ("Proposed Solution", req.proposed_solution or ""),
            ("Confidence Score", f"{int((req.confidence_score or 0.75) * 100)}%"),
        ]
        curr_row = 3
        for label, val in rows:
            ws1.cell(row=curr_row, column=1, value=label).font = s["bold_font"]
            ws1.cell(row=curr_row, column=1).border = s["border"]
            c = ws1.cell(row=curr_row, column=2, value=val)
            c.font = s["normal_font"]
            c.alignment = s["left"]
            c.border = s["border"]
            ws1.merge_cells(f"B{curr_row}:D{curr_row}")
            curr_row += 1
            
    ws1.column_dimensions["A"].width = 25
    ws1.column_dimensions["B"].width = 90

    # ── Sheet 2: Modules & Features ──────────────────────────────────────────
    if req:
        ws2 = wb.create_sheet(title="Requirements Analysis")
        _write_header_row(ws2, 1, ["Module", "Feature", "Priority", "Complexity", "Technical Description", "Acceptance Criteria"], s)
        feature_map = {f.name: f for f in (req.features or [])}
        row = 2
        if req.modules:
            for mod in req.modules:
                for fname in (mod.feature_names or []):
                    feat = feature_map.get(fname)
                    _write_data_row(ws2, row, [
                        mod.name,
                        fname,
                        feat.priority.upper() if feat else "",
                        feat.complexity.upper() if feat else "",
                        feat.description if feat else "",
                        "\n".join(feat.acceptance_criteria) if feat and feat.acceptance_criteria else "",
                    ], s, alt=(row % 2 == 0))
                    row += 1
        elif req.features:
            for feat in req.features:
                _write_data_row(ws2, row, [
                    "General",
                    feat.name,
                    feat.priority.upper(),
                    feat.complexity.upper(),
                    feat.description,
                    "\n".join(feat.acceptance_criteria)
                ], s, alt=(row % 2 == 0))
                row += 1
        ws2.column_dimensions["F"].width = 60

        # ── Sheet 3: Effort Estimation ──────────────────────────────────────────
        ws3 = wb.create_sheet(title="Effort Estimation")
        _write_header_row(ws3, 1, ["Module", "Feature", "Complexity", "Recommended Dev Hours", "Recommended Test Hours", "Total Hours"], s)
        row = 2
        req_del = req.delivery_plan if req.delivery_plan else None
        total_all_hrs = 0
        if req_del and req_del.feature_estimates:
            for est in req_del.feature_estimates:
                dev_hrs = sum(d.days * 8 for d in est.developer_days)
                test_hrs = sum(t.days * 8 for t in est.tester_days)
                _write_data_row(ws3, row, [
                    est.module_name or "General",
                    est.feature_name,
                    est.complexity.upper(),
                    int(round(dev_hrs)),
                    int(round(test_hrs)),
                    f"=ROUND(D{row}+E{row}, 0)"
                ], s, alt=(row % 2 == 0))
                row += 1
        elif req.features:
            for feat in req.features:
                _write_data_row(ws3, row, [
                    "General",
                    feat.name,
                    feat.complexity.upper(),
                    40, # Default dev hours
                    16, # Default test hours
                    f"=ROUND(D{row}+E{row}, 0)"
                ], s, alt=(row % 2 == 0))
                row += 1
            
        if row > 2:
            total_row = row + 1
            ws3.cell(row=total_row, column=1, value="TOTAL RECOMMENDED PROJECT HOURS").font = s["bold_font"]
            ws3.merge_cells(f"A{total_row}:E{total_row}")
            for col in range(1, 6):
                ws3.cell(row=total_row, column=col).border = s["border"]
            c_total = ws3.cell(row=total_row, column=6, value=f"=ROUND(SUM(F2:F{row-1}), 0)")
            c_total.font = s["bold_font"]
            c_total.alignment = s["center"]
            c_total.border = s["border"]

        ws3.column_dimensions["A"].width = 24
        ws3.column_dimensions["B"].width = 30
        ws3.column_dimensions["D"].width = 22
        ws3.column_dimensions["E"].width = 22
        ws3.column_dimensions["F"].width = 18

    wb.save(output)
    output.seek(0)
    filename = f"{data.title.replace(' ', '_')}_Estimator_SRS.xlsx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@router.post("/team")
async def export_team_excel(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User | None = Depends(get_optional_user)
):
    if current_user:
        history = UserHistory(
            user_id=current_user.id,
            action="Downloaded Team Excel",
            details=f"Exported Team Excel for: {data.get('project_name', 'Project')}"
        )
        db.add(history)
        db.commit()

    output = io.BytesIO()
    wb = Workbook()
    s = _make_styles()

    ws = wb.active
    ws.title = "Resource Allocation"

    ws.merge_cells("A1:C1")
    ws["A1"] = f"TEAM DESIGN: {data.get('project_name', 'Estimator Project')}"
    ws["A1"].font = s["title_font"]
    ws["A1"].fill = s["navy_fill"]
    ws["A1"].alignment = s["center"]
    ws.row_dimensions[1].height = 40

    _write_header_row(ws, 3, ["Position / Role", "Headcount", "Responsibilities & Justification"], s)
    
    members = data.get("members", [])
    row = 4
    for member in members:
        _write_data_row(ws, row, [
            member.get("role", "Consultant"),
            member.get("count", 0),
            member.get("description", "")
        ], s, alt=(row % 2 == 0), bold_first=True)
        row += 1

    total_row = row + 1
    ws.cell(row=total_row, column=1, value="TOTAL ESTIMATED HEADCOUNT").font = s["bold_font"]
    ws.cell(row=total_row, column=2, value=data.get("total_size", 0)).font = s["bold_font"]
    ws.cell(row=total_row, column=2).alignment = s["center"]
    
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 85

    wb.save(output)
    output.seek(0)
    filename = f"{data.get('project_name', 'Project').replace(' ', '_')}_Estimator_Team.xlsx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


def _team_project_name(data: dict) -> str:
    return data.get("project_name") or "Estimator Project"


def _team_members(data: dict) -> list[dict]:
    members = data.get("members") or []
    return [member for member in members if isinstance(member, dict)]


@router.post("/team/docx")
async def export_team_docx(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User | None = Depends(get_optional_user)
):
    project_name = _team_project_name(data)
    if current_user:
        history = UserHistory(
            user_id=current_user.id,
            action="Downloaded Team Word",
            details=f"Exported Team Word document for: {project_name}"
        )
        db.add(history)
        db.commit()

    document = Document()
    document.add_heading(f"Team Allocation: {project_name}", level=1)
    document.add_paragraph(f"Total Headcount: {data.get('total_size', 0)}")
    document.add_paragraph(f"Total Effort Hours: {round(float(data.get('total_project_hours') or 0), 2)}")
    document.add_paragraph(f"Project Duration: {round(float(data.get('total_working_weeks') or 0) * 40, 2)} hours")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Role", "People", "Hours / Person", "Active Weeks", "Responsibilities"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for member in _team_members(data):
        cells = table.add_row().cells
        cells[0].text = str(member.get("role") or "Consultant")
        cells[1].text = str(member.get("count") or 0)
        cells[2].text = str(member.get("hours_per_member") or 0)
        cells[3].text = str(member.get("active_weeks") or 0)
        cells[4].text = str(member.get("description") or "")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    filename = f"{project_name.replace(' ', '_')}_Estimator_Team.docx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("/team/pdf")
async def export_team_pdf(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User | None = Depends(get_optional_user)
):
    project_name = _team_project_name(data)
    if current_user:
        history = UserHistory(
            user_id=current_user.id,
            action="Downloaded Team PDF",
            details=f"Exported Team PDF for: {project_name}"
        )
        db.add(history)
        db.commit()

    output = io.BytesIO()
    doc = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(f"Team Allocation: {project_name}", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Total Headcount: {data.get('total_size', 0)}", styles["Normal"]),
        Paragraph(f"Total Effort Hours: {round(float(data.get('total_project_hours') or 0), 2)}", styles["Normal"]),
        Paragraph(f"Project Duration: {round(float(data.get('total_working_weeks') or 0) * 40, 2)} hours", styles["Normal"]),
        Spacer(1, 12),
    ]

    rows = [[
        Paragraph("<b>Role</b>", styles["Normal"]),
        Paragraph("<b>People</b>", styles["Normal"]),
        Paragraph("<b>Hours / Person</b>", styles["Normal"]),
        Paragraph("<b>Active Weeks</b>", styles["Normal"]),
        Paragraph("<b>Responsibilities</b>", styles["Normal"]),
    ]]
    for member in _team_members(data):
        rows.append([
            Paragraph(str(member.get("role") or "Consultant"), styles["Normal"]),
            Paragraph(str(member.get("count") or 0), styles["Normal"]),
            Paragraph(str(member.get("hours_per_member") or 0), styles["Normal"]),
            Paragraph(str(member.get("active_weeks") or 0), styles["Normal"]),
            Paragraph(str(member.get("description") or ""), styles["Normal"]),
        ])

    table = Table(rows, colWidths=[110, 44, 75, 75, 200], repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    doc.build(story)
    output.seek(0)

    filename = f"{project_name.replace(' ', '_')}_Estimator_Team.pdf"
    return Response(
        content=output.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@router.post("/bundle")
async def export_bundle(
    srs: SRSGenerationResult = Body(...),
    team: dict = Body(...),
    cost: CostEstimationExportRequest = Body(...),
    db: Session = Depends(get_db),
    current_user: User | None = Depends(get_optional_user)
):
    """Consolidated 4-Sheet Professional Excel Workbook."""
    if current_user:
        project_name_log = cost.project_name or team.get("project_name") or "Project"
        if srs.structured_requirements and srs.structured_requirements.project_name:
            project_name_log = srs.structured_requirements.project_name
        history = UserHistory(
            user_id=current_user.id,
            action="Downloaded Master Bundle",
            details=f"Exported Master Bundle for: {project_name_log}"
        )
        db.add(history)
        db.commit()

    output = io.BytesIO()
    wb = Workbook()
    s = _make_styles()
    project_name = (
        (srs.structured_requirements.project_name if srs.structured_requirements else None)
        or cost.project_name
        or team.get("project_name")
        or "Estimator Project"
    )
    totals = calculate_cost_totals(cost)
    total_weeks = determine_project_weeks(cost, team.get("total_working_weeks", 0) or 0)
    timeline_rows = build_module_timeline(
        srs.structured_requirements,
        totals["total_project_hours"],
        totals["grand_total"],
        total_weeks,
    )

    # ── Sheet 1: Readme (Introduction) ──────────────────────────────────────────
    ws_readme = wb.active
    ws_readme.title = "1. Readme"
    ws_readme.merge_cells("A1:C1")
    ws_readme["A1"] = f"PROJECT BUNDLE: {project_name}"
    ws_readme["A1"].font = s["title_font"]
    ws_readme["A1"].fill = s["navy_fill"]
    ws_readme["A1"].alignment = s["center"]
    ws_readme.row_dimensions[1].height = 40

    _write_header_row(ws_readme, 3, ["Document Tab", "Description of Contents", "Data Status"], s)
    contents = [
        ("Approved SRS", "Detailed functional requirements, modules, and extraction notes.", "FINALIZED"),
        ("Team & Effort", "Resource mix, role breakdown, and feature-wise effort estimation (hours).", "APPROVED"),
        ("Cost Estimation", "Budget breakdown, hourly rates, and total project valuation.", "DRAFT"),
    ]
    for i, row in enumerate(contents, 4):
        _write_data_row(ws_readme, i, row, s, alt=(i % 2 == 0))
    
    ws_readme.column_dimensions["A"].width = 25
    ws_readme.column_dimensions["B"].width = 65
    ws_readme.column_dimensions["C"].width = 20

    # ── Sheet 2: SRS Sections (identical content to the UI review view) ──────────
    ws_srs = wb.create_sheet("2. SRS")
    ws_srs.merge_cells("A1:B1")
    ws_srs["A1"] = f"SRS REQUIREMENTS: {project_name}"
    ws_srs["A1"].font = s["title_font"]
    ws_srs["A1"].fill = s["navy_fill"]
    ws_srs["A1"].alignment = s["center"]
    ws_srs.row_dimensions[1].height = 40

    _write_header_row(ws_srs, 3, ["Section Title", "Section Content"], s)
    srs_row = 4
    for section in srs.sections:
        body = section.body
        if "team design and working hours" in section.title.lower():
            body = body + build_synced_team_section(project_name, cost, totals["total_project_hours"], total_weeks)
        c_title = ws_srs.cell(row=srs_row, column=1, value=section.title)
        c_title.font = s["bold_font"]
        c_title.border = s["border"]
        c_title.alignment = s["left"]
        c_body = ws_srs.cell(row=srs_row, column=2, value=body)
        c_body.font = s["normal_font"]
        c_body.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c_body.border = s["border"]
        line_count = max(1, body.count("\n") + 1)
        ws_srs.row_dimensions[srs_row].height = min(400, max(30, line_count * 15))
        srs_row += 1
    ws_srs.column_dimensions["A"].width = 35
    ws_srs.column_dimensions["B"].width = 120

    # ── Sheet 3: Team Allocation ───────────────────────────────────────────────
    ws_team = wb.create_sheet("3. Team Allocation")
    ws_team.merge_cells("A1:F1")
    ws_team["A1"] = f"TEAM ALLOCATION: {project_name}"
    ws_team["A1"].font = s["title_font"]
    ws_team["A1"].fill = s["navy_fill"]
    ws_team["A1"].alignment = s["center"]
    ws_team.row_dimensions[1].height = 40

    _write_header_row(ws_team, 3, ["Role", "Headcount", "Weekly Hours", "Hours / Member", "Total Hours", "Justification / Role Notes"], s)
    row = 4
    team_notes = {str(member.get("role", "")).lower(): member.get("description", "") for member in team.get("members", [])}
    for member in cost.members:
        _write_data_row(ws_team, row, [
            member.role,
            int(round(safe_float(member.count))),
            int(round(safe_float(member.weekly_hours))),
            int(round(safe_float(member.hours_per_member))),
            f"=ROUND(B{row}*D{row}, 0)",
            team_notes.get(member.role.lower(), "")
        ], s, alt=(row % 2 == 0), bold_first=True)
        row += 1
    
    ws_team.cell(row=row + 1, column=1, value="TOTAL PROJECT HOURS").font = s["bold_font"]
    c_tot_hrs = ws_team.cell(row=row + 1, column=5, value=f"=ROUND(SUM(E4:E{row-1}), 0)")
    c_tot_hrs.font = s["bold_font"]
    c_tot_hrs.border = s["border"]
    c_tot_hrs.alignment = s["center"]
    
    ws_team.cell(row=row + 2, column=1, value="TOTAL WORKING WEEKS").font = s["bold_font"]
    c_tot_wks = ws_team.cell(row=row + 2, column=3, value=int(round(total_weeks)))
    c_tot_wks.font = s["bold_font"]
    c_tot_wks.border = s["border"]
    c_tot_wks.alignment = s["center"]

    ws_team.column_dimensions["A"].width = 30
    ws_team.column_dimensions["B"].width = 15
    ws_team.column_dimensions["C"].width = 15
    ws_team.column_dimensions["D"].width = 18
    ws_team.column_dimensions["E"].width = 18
    ws_team.column_dimensions["F"].width = 70

    # ── Sheet 4: Cost Estimation (Financials) ───────────────────────────────────
    ws_cost = wb.create_sheet("4. Cost Estimation")
    ws_cost.merge_cells("A1:G1")
    ws_cost["A1"] = f"COST ESTIMATION: {project_name}"
    ws_cost["A1"].font = s["title_font"]
    ws_cost["A1"].fill = s["navy_fill"]
    ws_cost["A1"].alignment = s["center"]
    ws_cost.row_dimensions[1].height = 40

    _write_header_row(
        ws_cost,
        3,
        ["Role", "Count", f"Hourly Rate ({cost.currency})", "Weekly Hours", "Hours / Member", f"Cost / Employee ({cost.currency})", f"Total ({cost.currency})"],
        s,
    )

    dev_rows = []
    test_rows = []
    deploy_rows = []

    row = 4
    for item in member_breakdown:
        role_name = item["role"]
        if not is_management_role(role_name):
            if is_testing_role(role_name):
                test_rows.append(row)
            elif is_deployment_role(role_name):
                deploy_rows.append(row)
            else:
                dev_rows.append(row)

        values = [
            role_name,
            item["count"],
            item["hourly_rate"],
            item["weekly_hours"],
            item["hours_per_member"],
            f"=C{row}*E{row}",
            f"=B{row}*F{row}",
        ]
        for col_idx, val in enumerate(values, 1):
            c = ws_cost.cell(row=row, column=col_idx, value=val)
            c.font = s["normal_font"]
            c.border = s["border"]
            c.alignment = s["left"] if col_idx == 1 else s["center"]
            if col_idx in (3, 6, 7):
                c.number_format = f'"{cost.currency}" #,##0'
        row += 1

    member_end_row = row - 1

    salary_total = totals.get("salary_total", 0)
    pm_cost = totals.get("project_management", 0)
    pm_ratio = (pm_cost / salary_total) if salary_total > 0 else 0.10

    effort_subtotal = salary_total + pm_cost + totals.get("misc_total", 0)
    risk_cost = totals.get("risk_contingency", 0)
    risk_ratio = (risk_cost / effort_subtotal) if effort_subtotal > 0 else 0.10

    nego_cost = totals.get("negotiation_buffer", 0)
    nego_ratio = (nego_cost / effort_subtotal) if effort_subtotal > 0 else 0.05

    summary_start = row + 2
    summary_rows = [
        ("Development Total Cost", f"=ROUND(SUM({','.join(f'G{r}' for r in dev_rows)}), 0)" if dev_rows else "=0"),
        ("Testing Total Cost", f"=ROUND(SUM({','.join(f'G{r}' for r in test_rows)}), 0)" if test_rows else "=0"),
        ("Deployment Total Cost", f"=ROUND(SUM({','.join(f'G{r}' for r in deploy_rows)}), 0)" if deploy_rows else "=0"),
        ("Team Salary Total", f"=ROUND(SUM(G4:G{member_end_row}), 0)"),
        ("Project Management", f"=ROUND(G{summary_start+3}*{pm_ratio}, 0)"),
        ("Miscellaneous Costs", int(round(totals.get("misc_total", 0)))),
        ("Risk Contingency", f"=ROUND((G{summary_start+3}+G{summary_start+4}+G{summary_start+5})*{risk_ratio}, 0)"),
        ("Negotiation Buffer", f"=ROUND((G{summary_start+3}+G{summary_start+4}+G{summary_start+5})*{nego_ratio}, 0)"),
        ("Project Total Cost Estimation", f"=ROUND(G{summary_start+3}+G{summary_start+4}+G{summary_start+5}+G{summary_start+6}+G{summary_start+7}, 0)"),
        ("Profit Slabs", int(round(totals.get("profit_total", 0)))),
    ]

    for offset, (label, formula_or_value) in enumerate(summary_rows):
        curr_row = summary_start + offset
        ws_cost.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=6)
        label_cell = ws_cost.cell(row=curr_row, column=1, value=label)
        value_cell = ws_cost.cell(row=curr_row, column=7, value=formula_or_value)
        
        label_cell.font = s["bold_font"]
        value_cell.font = s["bold_font"]
        label_cell.border = s["border"]
        value_cell.border = s["border"]
        label_cell.alignment = s["left"]
        value_cell.alignment = s["center"]
        value_cell.number_format = f'"{cost.currency}" #,##0'
        for col_idx in range(2, 7):
            ws_cost.cell(row=curr_row, column=col_idx).border = s["border"]

    grand_total_row = summary_start + len(summary_rows) + 1
    ws_cost.merge_cells(start_row=grand_total_row, start_column=1, end_row=grand_total_row, end_column=6)
    label_cell = ws_cost.cell(row=grand_total_row, column=1, value="GRAND TOTAL ESTIMATION")
    value_cell = ws_cost.cell(row=grand_total_row, column=7, value=f"=ROUND(G{summary_start+8}+G{summary_start+9}, 0)")
    
    label_cell.font = s["bold_font"]
    value_cell.font = Font(name="Calibri", size=12, bold=True, color="000000")
    label_cell.border = s["border"]
    value_cell.border = s["border"]
    label_cell.alignment = s["left"]
    value_cell.alignment = s["center"]
    value_cell.number_format = f'"{cost.currency}" #,##0'
    for col_idx in range(2, 7):
        ws_cost.cell(row=grand_total_row, column=col_idx).border = s["border"]

    ws_cost.column_dimensions["A"].width = 30
    ws_cost.column_dimensions["B"].width = 15
    ws_cost.column_dimensions["C"].width = 18
    ws_cost.column_dimensions["D"].width = 15
    ws_cost.column_dimensions["E"].width = 18
    ws_cost.column_dimensions["F"].width = 20
    ws_cost.column_dimensions["G"].width = 25

    # ── Sheet 5: Gantt Chart (Timeline & Cost) ──────────────────────────────────
    ws_gantt = wb.create_sheet("5. Gantt Chart")
    week_headers = [f"Wk {index}" for index in range(1, max(6, total_weeks) + 1)]
    headers = ["Module / Phase", "Estimated Hours", f"Allocated Cost ({cost.currency})"] + week_headers
    ws_gantt.merge_cells(f"A1:{get_column_letter(len(headers))}1")
    ws_gantt["A1"] = f"PROJECT GANTT & TIMELINE: {project_name}"
    ws_gantt["A1"].font = s["title_font"]
    ws_gantt["A1"].fill = s["navy_fill"]
    ws_gantt["A1"].alignment = s["center"]
    ws_gantt.row_dimensions[1].height = 40

    _write_header_row(ws_gantt, 3, headers, s)
    
    g_row = 4
    for entry in timeline_rows:
        vals = [entry["label"], int(round(entry["hours"])), int(round(entry["cost"]))] + entry["markers"]
        _write_data_row(ws_gantt, g_row, vals, s, alt=(g_row % 2 == 0))
        ws_gantt.cell(row=g_row, column=3).number_format = f'"{cost.currency}" #,##0'
        g_row += 1
            
    # Summary Info after Gantt
    g_row += 2
    
    ws_gantt.cell(row=g_row, column=1, value="TOTAL PROJECT TIME (HOURS)").font = s["bold_font"]
    ws_gantt.cell(row=g_row, column=2, value=f"=ROUND(SUM(B4:B{g_row-3}), 0)").font = s["normal_font"]
    g_row += 1
    
    ws_gantt.cell(row=g_row, column=1, value="TOTAL WORKING WEEKS").font = s["bold_font"]
    ws_gantt.cell(row=g_row, column=2, value=int(round(total_weeks))).font = s["normal_font"]
    g_row += 1
    
    ws_gantt.cell(row=g_row, column=1, value="GRAND TOTAL COST").font = s["bold_font"]
    c_gantt_total = ws_gantt.cell(row=g_row, column=2, value=f"=ROUND(SUM(C4:C{g_row-5}), 0)")
    c_gantt_total.font = s["normal_font"]
    c_gantt_total.number_format = f'"{cost.currency}" #,##0'
    g_row += 1

    for column_index in range(1, len(headers) + 1):
        col_let = get_column_letter(column_index)
        ws_gantt.column_dimensions[col_let].width = 30 if col_let == "A" else (25 if col_let == "C" else 15)

    wb.save(output)
    output.seek(0)
    
    filename = f"{project_name.replace(' ', '_')}_Estimator_Master.xlsx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
