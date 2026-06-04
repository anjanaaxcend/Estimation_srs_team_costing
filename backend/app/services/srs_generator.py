from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import settings
from app.schemas.client import ClientInput
from app.schemas.requirements import RequirementExtractionResult
from app.schemas.srs import AppendixItem, DeliveryPlanData, ModelSelection, ReferenceItem, SRSGenerationResult, SRSSection
from app.services.openai_srs_generator import OpenAISRSGenerator
from app.services.rag_service import RAGService
from app.utils.slug import slugify
from app.utils.project_name import resolve_project_name

# Pre-fetched real-time industry data for 2026 (Jewellery & E-commerce)
INDUSTRY_INSIGHTS_2026 = """
In 2026, the architecture of successful jewellery e-commerce is defined by unified commerce.
Key tech standards:
- API-First & Headless/Composable Commerce.
- Real-Time Inventory Synchronization (RapNet/Nivoda APIs).
- Augmented Reality (AR) Virtual Try-On (VTO) for rings and necklaces.
- AI-Powered Hyper-personalization and Agentic Commerce.
- High-definition 360-degree views and mobile-first performance engineering.
"""

class SRSGeneratorService:
    def __init__(self) -> None:
        self.template = json.loads(settings.srs_template_path.read_text(encoding="utf-8"))
        settings.generated_dir.mkdir(parents=True, exist_ok=True)
        self.openai_generator = OpenAISRSGenerator()
        self.rag_service = RAGService()

    def generate(
        self,
        client_input: ClientInput,
        requirements: RequirementExtractionResult,
        references: list[ReferenceItem] | None = None,
        appendices: list[AppendixItem] | None = None,
        rag_context: list[str] | None = None,
        web_context: str | None = None,
        selected_model: ModelSelection | None = None,
        prior_sections: list[SRSSection] | None = None,
        regeneration_feedback: str = "",
    ) -> SRSGenerationResult:
        resolved_project_name = resolve_project_name(
            extracted_name=requirements.project_name,
            provided_name=client_input.project_name,
            raw_text=requirements.normalized_text or client_input.raw_text,
        )
        resolved_client_input = client_input.model_copy(update={"project_name": resolved_project_name})
        resolved_requirements = requirements.model_copy(update={"project_name": resolved_project_name})
        section_titles = list(self.template["sections"])

        if rag_context is None:
            rag_context = self.rag_service.query(resolved_requirements.normalized_text or resolved_client_input.raw_text)

        if web_context is None:
            web_context = INDUSTRY_INSIGHTS_2026

        sections, delivery_plan, resolved_model = self._generate_sections(
            client_input=resolved_client_input,
            requirements=resolved_requirements,
            references=references or [],
            appendices=appendices or [],
            section_titles=section_titles,
            rag_context=rag_context,
            web_context=web_context,
            selected_model=selected_model,
            prior_sections=prior_sections or [],
            regeneration_feedback=regeneration_feedback,
        )

        filename_root = slugify(resolved_project_name)
        xlsx_path = settings.generated_dir / f"{filename_root}.xlsx"
        docx_path = settings.generated_dir / f"{filename_root}.docx"
        pdf_path = settings.generated_dir / f"{filename_root}.pdf"

        self._build_xlsx(resolved_project_name, resolved_requirements, sections, delivery_plan, xlsx_path)
        self._build_docx(resolved_project_name, sections, docx_path)
        self._build_pdf(resolved_project_name, sections, pdf_path)

        return SRSGenerationResult(
            title=f"{resolved_project_name} - SRS Software Requirements Specifications",
            sections=sections,
            delivery_plan=delivery_plan,
            docx_path=f"/generated/{docx_path.name}",
            xlsx_path=f"/generated/{xlsx_path.name}",
            pdf_path=f"/generated/{pdf_path.name}",
            cleaned_text=resolved_requirements.normalized_text or resolved_client_input.raw_text,
            structured_requirements=resolved_requirements,
            selected_model=resolved_model,
        )

    def _generate_sections(
        self,
        client_input: ClientInput,
        requirements: RequirementExtractionResult,
        references: list[ReferenceItem],
        appendices: list[AppendixItem],
        section_titles: list[str],
        rag_context: list[str] | None = None,
        web_context: str | None = None,
        selected_model: ModelSelection | None = None,
        prior_sections: list[SRSSection] | None = None,
        regeneration_feedback: str = "",
    ) -> tuple[list[SRSSection], DeliveryPlanData | None, ModelSelection]:
        if not self.openai_generator.is_enabled(selected_model):
            raise RuntimeError(
                "No configured backend SRS generation model is available. Update backend/.env with a supported provider key and model."
            )

        return self.openai_generator.generate_sections(
            client_input=client_input,
            requirements=requirements,
            references=references,
            appendices=appendices,
            section_titles=section_titles,
            rag_context=rag_context,
            web_context=web_context,
            selected_model=selected_model,
            prior_sections=prior_sections or [],
            regeneration_feedback=regeneration_feedback,
        )

    def _build_section_map(
        self,
        client_input: ClientInput,
        requirements: RequirementExtractionResult,
        references: list[ReferenceItem],
        appendices: list[AppendixItem],
    ) -> dict[str, str]:
        introduction = self._build_introduction(client_input, requirements)
        overall_description = self._build_overall_description(client_input, requirements)
        functional_requirements = self._build_system_features_and_functional_requirements(requirements)
        external_interfaces = self._build_external_interface_requirements(client_input, requirements)
        non_functional_requirements = self._build_non_functional_requirements(requirements)
        technologies_and_constraints = self._build_technologies_and_design_constraints(requirements)
        delivery_plan = self._build_project_delivery_plan(requirements)
        appendices_and_references = self._build_appendices_and_references(requirements, references, appendices)

        return {
            "1. Introduction": introduction,
            "2. Project Modules and Features": functional_requirements,
            "3. UI Pages and Screen Design": "UI pages and screen design details will be generated based on features.",
            "4. External Interface Requirements": external_interfaces,
            "5. Non-Functional Requirements": non_functional_requirements,
            "6. Technologies and Design Constraints": technologies_and_constraints,
            "7. Team Design and Working Hours": delivery_plan,
            # Legacy title compatibility for older templates.
            "Problem Overview": introduction,
            "Client Requirements and Objectives": overall_description,
            "Proposed Solution": self._build_proposed_solution(requirements),
            "Modules and Feature Breakdown": functional_requirements,
            "Roles and Delivery Plan": delivery_plan,
            "Technologies and Tools": technologies_and_constraints,
            "Conclusion": appendices_and_references,
        }

    def _build_introduction(self, client_input: ClientInput, requirements: RequirementExtractionResult) -> str:
        lines = [
            "Purpose",
            f"- This document defines the software requirements for {client_input.project_name}.",
            "",
            "Project Context",
            f"- Client Name: {client_input.client_name or 'Not specified'}",
            f"- Industry: {client_input.industry or 'Not specified'}",
            "",
            "Problem Statement",
            requirements.problem_statement or requirements.executive_summary or "Problem statement was not explicitly provided.",
        ]
        return "\n".join(lines)

    def _build_overall_description(self, client_input: ClientInput, requirements: RequirementExtractionResult) -> str:
        goals = requirements.project_objectives or client_input.business_goals
        objective_lines = [f"- {goal}" for goal in goals] or ["- No explicit objectives were identified in the uploaded brief."]
        assumption_lines = [f"- {item}" for item in requirements.assumptions] or ["- No additional assumptions were captured."]
        lines = [
            "Product Perspective",
            requirements.proposed_solution or requirements.executive_summary or "High-level solution perspective was not provided.",
            "",
            "Project Objectives",
            *objective_lines,
            "",
            "Assumptions",
            *assumption_lines,
        ]
        return "\n".join(lines)

    def _build_system_features_and_functional_requirements(self, requirements: RequirementExtractionResult) -> str:
        lines = ["Functional Requirements by Module"]
        feature_lookup = {feature.name: feature for feature in requirements.features}
        if requirements.modules:
            for module in requirements.modules:
                lines.append(f"Module: {module.name}")
                lines.append(module.summary or "No module summary supplied.")
                if module.feature_names:
                    for feature_name in module.feature_names:
                        feature = feature_lookup.get(feature_name)
                        if feature is None:
                            lines.append(f"- {feature_name}")
                            continue
                        lines.append(f"- {feature.name} | Priority: {feature.priority} | Complexity: {feature.complexity}")
                        lines.append(f"  Description: {feature.description}")
                        if feature.acceptance_criteria:
                            lines.append("  Acceptance Criteria")
                            lines.extend(f"  - {criterion}" for criterion in feature.acceptance_criteria)
                else:
                    lines.append("- No features were mapped to this module.")
                lines.append("")
        elif requirements.features:
            for feature in requirements.features:
                lines.append(f"- {feature.name} | Priority: {feature.priority} | Complexity: {feature.complexity}")
                lines.append(f"  Description: {feature.description}")
                if feature.acceptance_criteria:
                    lines.extend(f"  - {criterion}" for criterion in feature.acceptance_criteria)
        else:
            lines.append("- No functional requirements were inferred from the uploaded brief.")

        return "\n".join(lines).rstrip()

    def _build_external_interface_requirements(self, client_input: ClientInput, requirements: RequirementExtractionResult) -> str:
        integration_lines = [f"- {item}" for item in client_input.integrations] or ["- No external integrations were specified."]
        deployment_lines = [f"- {item}" for item in client_input.deployment_preferences] or ["- Deployment interfaces were not specified."]
        data_model_lines = [
            f"- {model.name}: {model.description} | Attributes: {', '.join(model.attributes) or 'None'}"
            for model in requirements.data_models
        ] or ["- No data models were identified."]
        backend_job_lines = [
            f"- {job.name} | Trigger: {job.trigger_type} | Description: {job.description}"
            for job in requirements.backend_jobs
        ] or ["- No backend jobs were identified."]
        lines = [
            "Third-Party Integrations",
            *integration_lines,
            "",
            "Deployment and Hosting Interfaces",
            *deployment_lines,
            "",
            "Data Interfaces",
            *data_model_lines,
            "",
            "Background Processing Interfaces",
            *backend_job_lines,
        ]
        return "\n".join(lines)

    def _build_non_functional_requirements(self, requirements: RequirementExtractionResult) -> str:
        nfr_lines = [
            f"- {item.category}: {item.description} | Target: {item.measurable_target}"
            for item in requirements.non_functional_requirements
        ] or ["- No non-functional requirements were identified."]
        lines = [
            "Quality Attributes and Targets",
            *nfr_lines,
            "",
            "Extraction Confidence",
            f"- Confidence Score: {requirements.confidence_score}",
        ]
        return "\n".join(lines)

    def _build_technologies_and_design_constraints(self, requirements: RequirementExtractionResult) -> str:
        technology_lines = [f"- {item}" for item in requirements.recommended_technologies] or ["- No explicit technology recommendations were derived."]
        tool_lines = [f"- {item}" for item in requirements.recommended_tools] or ["- No explicit tool recommendations were derived."]
        constraint_lines = [f"- {item.category}: {item.description}" for item in requirements.constraints] or ["- No explicit constraints were identified."]
        lines = [
            "Recommended Technologies",
            *technology_lines,
            "",
            "Recommended Tools",
            *tool_lines,
            "",
            "Design and Implementation Constraints",
            *constraint_lines,
        ]
        return "\n".join(lines)

    def _build_project_delivery_plan(self, requirements: RequirementExtractionResult) -> str:
        return self._build_roles_and_delivery_plan(requirements)

    def _build_appendices_and_references(
        self,
        requirements: RequirementExtractionResult,
        references: list[ReferenceItem],
        appendices: list[AppendixItem],
    ) -> str:
        return self._build_conclusion(requirements, references, appendices)

    def _build_problem_overview(self, client_input: ClientInput, requirements: RequirementExtractionResult) -> str:
        lines = [
            f"Project Name: {client_input.project_name}",
            f"Client Name: {client_input.client_name or 'Not specified'}",
            f"Industry: {client_input.industry or 'Not specified'}",
            "",
            "Problem Overview",
            requirements.problem_statement or requirements.executive_summary,
            "",
            "Normalized Source Text",
            requirements.normalized_text or client_input.raw_text,
        ]
        return "\n".join(lines)

    def _build_client_requirements_and_objectives(self, client_input: ClientInput, requirements: RequirementExtractionResult) -> str:
        goals = requirements.project_objectives or client_input.business_goals
        assumption_lines = [f"- {item}" for item in requirements.assumptions] or ["- No additional assumptions were captured."]
        objective_lines = [f"- {goal}" for goal in goals] if goals else ["- No explicit objectives were identified in the uploaded brief."]
        constraint_lines = [f"- {item.category}: {item.description}" for item in requirements.constraints] or ["- No explicit constraints were identified."]
        lines = [
            "Client Objectives",
            *objective_lines,
            "",
            "Known Constraints",
            *constraint_lines,
            "",
            "Assumptions",
            *assumption_lines,
        ]
        return "\n".join(lines)

    def _build_proposed_solution(self, requirements: RequirementExtractionResult) -> str:
        constraint_lines = [f"- {item.category}: {item.description}" for item in requirements.constraints]
        lines = [
            requirements.proposed_solution or "A proposed company solution was not identified from the source material.",
        ]
        if constraint_lines:
            lines.extend(["", "Solution Constraints", *constraint_lines])
        return "\n".join(lines)

    def _build_modules_and_feature_breakdown(self, requirements: RequirementExtractionResult) -> str:
        lines = ["Project Modules"]
        if requirements.modules:
            feature_lookup = {feature.name: feature for feature in requirements.features}
            for module in requirements.modules:
                lines.append(f"Module: {module.name}")
                lines.append(module.summary)
                if module.feature_names:
                    for feature_name in module.feature_names:
                        feature = feature_lookup.get(feature_name)
                        if feature is None:
                            lines.append(f"- {feature_name}")
                            continue
                        lines.append(
                            f"- {feature.name} | Priority: {feature.priority} | Complexity: {feature.complexity}"
                        )
                        lines.append(f"  {feature.description}")
                else:
                    lines.append("- No features were mapped to this module.")
                lines.append("")
        else:
            lines.append("No modules were inferred from the uploaded brief.")

        return "\n".join(lines).rstrip()

    def _build_roles_and_delivery_plan(self, requirements: RequirementExtractionResult) -> str:
        role_lines = [
            f"- {role.name}: {', '.join(role.responsibilities)}"
            for role in requirements.user_roles
        ] or ["- No user roles were provided in the input."]
        delivery = requirements.delivery_plan
        delivery_lines = [
            f"Team Size: {delivery.team_size}",
            f"Developers: {delivery.developer_count}",
            f"Testers: {delivery.tester_count}",
            f"Developer Ratio: {delivery.developer_time_ratio}",
            f"Tester Ratio: {delivery.tester_time_ratio}",
            "",
            "Estimated Project Hours",
            *[f"- {item.level}: {round((item.days) * 8)} hours" for item in delivery.estimated_project_days],
            "",
            "Feature Delivery Allocation (Hours)",
        ]
        for estimate in delivery.feature_estimates:
            delivery_lines.append(
                f"- {estimate.feature_name} ({estimate.module_name}) | Recommended Dev: {estimate.recommended_developer_level} | Recommended Test: {estimate.recommended_tester_level}"
            )
            delivery_lines.extend(
                f"  Dev {item.level}: {round(item.days * 8)}h" for item in estimate.developer_days
            )
            delivery_lines.extend(
                f"  Test {item.level}: {round(item.days * 8)}h" for item in estimate.tester_days
            )
        if delivery.planning_assumptions:
            delivery_lines.extend(["", "Planning Assumptions"])
            delivery_lines.extend(f"- {item}" for item in delivery.planning_assumptions)

        return "\n".join(["Primary Roles", *role_lines, "", *delivery_lines]).rstrip()

    def _build_technologies_and_tools(self, requirements: RequirementExtractionResult) -> str:
        technology_lines = [f"- {item}" for item in requirements.recommended_technologies]
        tool_lines = [f"- {item}" for item in requirements.recommended_tools]
        nfr_lines = [
            f"- {item.category}: {item.description} | Target: {item.measurable_target}"
            for item in requirements.non_functional_requirements
        ]
        lines = [
            "Recommended Technologies",
            *(technology_lines or ["- No explicit technology recommendations were derived."]),
            "",
            "Recommended Tools",
            *(tool_lines or ["- No explicit tool recommendations were derived."]),
            "",
            "Quality and Non-Functional Requirements",
            *(nfr_lines or ["- No non-functional requirements were identified."]),
        ]
        return "\n".join(lines)

    def _build_conclusion(
        self,
        requirements: RequirementExtractionResult,
        references: list[ReferenceItem],
        appendices: list[AppendixItem],
    ) -> str:
        lines = [requirements.conclusion or "The extracted requirements are ready for review and implementation planning."]
        if requirements.ai_observations:
            lines.extend(["", "Analyst Notes"])
            lines.extend(f"- {item}" for item in requirements.ai_observations)
        if references:
            lines.extend(["", "References"])
            lines.extend(f"- {reference.title}: {reference.description}" for reference in references)
        if appendices:
            lines.extend(["", "Appendices"])
            lines.extend(f"- {appendix.title}: {appendix.content}" for appendix in appendices)
        return "\n".join(lines)

    def _build_docx(self, project_name: str, sections: list[SRSSection], target_path: Path) -> None:
        document = Document()
        document.add_heading(f"{project_name} - Software Requirements Specification", level=0)
        document.add_paragraph("Prepared from the uploaded client brief using structured requirement extraction and SRS generation.")

        for section in sections:
            document.add_heading(section.title, level=1)
            for line in section.body.split("\n"):
                stripped = line.strip()
                if not stripped:
                    document.add_paragraph("")
                elif stripped.startswith("- "):
                    document.add_paragraph(stripped[2:], style="List Bullet")
                else:
                    document.add_paragraph(stripped)

        document.save(target_path)

    def _build_xlsx(
        self,
        project_name: str,
        requirements: RequirementExtractionResult,
        sections: list[SRSSection],
        delivery_plan: DeliveryPlanData | None,
        target_path: Path,
    ) -> None:
        workbook = Workbook()
        
        # Sheet 1: SRS Sections
        sheet = workbook.active
        sheet.title = "SRS Sections"
        heading_font = Font(size=14, bold=True)
        content_font = Font(size=12)
        justified_alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)

        sheet.append(["Section Name", "Content"])
        sheet["A1"].font = heading_font
        sheet["B1"].font = heading_font
        sheet.column_dimensions["A"].width = 42
        sheet.column_dimensions["B"].width = 120

        title_body = f"{project_name} - SRS Software Requirements Specifications"
        sheet.append(["Document Title", title_body])
        for section in sections:
            sheet.append([section.title, section.body])

        for row_index in range(2, sheet.max_row + 1):
            sheet.cell(row=row_index, column=1).font = heading_font
            sheet.cell(row=row_index, column=2).alignment = justified_alignment

        # Sheet 2: Delivery & Gantt
        if delivery_plan:
            gantt = workbook.create_sheet("Delivery Timeline")
            from openpyxl.styles import PatternFill
            
            # Styles
            header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            module_fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
            gantt_fill = PatternFill(start_color="0078D4", end_color="0078D4", fill_type="solid")

            # Team Recommendation Header
            gantt.merge_cells("A1:E1")
            gantt["A1"] = "RECOMMENDED TEAM STRUCTURE"
            gantt["A1"].font = heading_font

            team = delivery_plan.recommended_team
            if team:
                gantt.append(["Role", "Count"])
                gantt.append(["Lead Developers", team.lead_count])
                gantt.append(["Mid-level Developers", team.mid_count])
                gantt.append(["Junior Developers", team.junior_count])
                gantt.append(["QA Testers", team.tester_count])
                gantt.append(["Total Team Members", team.lead_count + team.mid_count + team.junior_count + team.tester_count])

            gantt.append([])
            gantt.append(["TOTAL ESTIMATED PROJECT DURATION", f"{delivery_plan.total_duration_days} days"])

            gantt.append([])
            gantt.append(["VISUAL DELIVERY TIMELINE (GANTT)"])

            # Module Headers
            cols = ["Module", "Estimated Dev Days", "Testing Days"]
            # Add week numbers as columns
            max_weeks = int(max((m.end_week for m in delivery_plan.modules), default=12)) + 1
            for w in range(1, max_weeks + 1):
                cols.append(f"W{w}")

            gantt.append(cols)
            header_row = gantt.max_row
            for c in range(1, len(cols) + 1):
                cell = gantt.cell(row=header_row, column=c)
                cell.fill = header_fill
                cell.font = header_font

            for mod in delivery_plan.modules:
                row_data = [
                    mod.module_name,
                    f"{mod.total_days}d",
                    f"{mod.testing_days}d",
                ]
                # Placeholder for Gantt bars
                row_data.extend([""] * max_weeks)
                gantt.append(row_data)
                curr_row = gantt.max_row

                # Fill Gantt blocks
                start_col = 4 + int(mod.start_week) - 1
                end_col = 4 + int(mod.end_week) - 1
                for c in range(start_col, end_col + 1):
                    if c <= gantt.max_column:
                        cell = gantt.cell(row=curr_row, column=c)
                        cell.fill = gantt_fill

            gantt.column_dimensions["A"].width = 30

        workbook.save(target_path)

    def _build_pdf(self, project_name: str, sections: list[SRSSection], target_path: Path) -> None:
        styles = getSampleStyleSheet()
        story = [
            Paragraph(f"{project_name} - Software Requirements Specification", styles["Title"]),
            Spacer(1, 18),
        ]

        for section in sections:
            story.append(Paragraph(section.title, styles["Heading2"]))
            for line in section.body.split("\n"):
                story.append(Paragraph(line if line.strip() else " ", styles["BodyText"]))
            story.append(Spacer(1, 12))

        document = SimpleDocTemplate(str(target_path))
        document.build(story)
